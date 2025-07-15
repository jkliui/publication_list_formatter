import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 要加粗下划线的作者字符串（精准匹配）
target_author = "Xxx, X"

def format_authors(authors_str):
    authors = authors_str.split(';')
    formatted = []
    for author in authors:
        author = author.strip()
        if ',' in author:
            last, first = author.split(',', 1)
            first_initial = first.strip()[0] if first.strip() else ''
            formatted.append(f"{last.strip()}, {first_initial}.")
        else:
            formatted.append(author.strip())  # fallback
    return '; '.join(formatted)

# 读取 CSV 文件
df = pd.read_csv("publication_list.csv")

# 替换 NaN 为 '' 并类型转换
df.fillna('', inplace=True)
df['Volume'] = pd.to_numeric(df['Volume'], errors='coerce').fillna(0).astype(int)
df['Year'] = pd.to_numeric(df['Year'], errors='coerce').fillna(0).astype(int)
df = df.astype({col: str for col in df.columns if col not in ['Volume', 'Year']})

# 格式化作者名
df['Formatted Authors'] = df['Authors'].apply(format_authors)

# 提取字段并排序（按照年份降序）
df_selected = df[['Formatted Authors', 'Title', 'Publication', 'Year', 'Volume', 'Pages']]
df_selected = df_selected.sort_values(by='Year', ascending=False).reset_index(drop=True)

# 创建 Word 文档
doc = Document()

# 设置默认字体样式
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# 总文章数
total = len(df_selected)

# 遍历每一篇文章
for _, row in df_selected.iterrows():
    author = row['Formatted Authors']
    title = row['Title']
    journal = row['Publication']
    year = row['Year']
    volume = row['Volume']
    pages = row['Pages']

    # 添加段落并设置格式
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(-0.3)

    # 写入编号
    paragraph.add_run(f"{total})  ")

    # 写入作者（处理 target_author 是否存在）
    if target_author in author:
        parts = author.split(target_author)
        for i, part in enumerate(parts):
            if part:
                paragraph.add_run(part)
            if i < len(parts) - 1:
                run = paragraph.add_run(target_author)
                run.bold = True
                run.underline = True
    else:
        paragraph.add_run(author)

    paragraph.add_run(" ")  # 加空格

    # Title
    paragraph.add_run(f"{title}. ")

    # Journal name (italic)
    paragraph.add_run(journal).italic = True
    paragraph.add_run(". ")

    # Year (bold)
    year_run = paragraph.add_run(str(year))
    year_run.bold = True
    paragraph.add_run(", ")

    # Volume (italic)
    volume_run = paragraph.add_run(str(volume))
    volume_run.italic = True
    paragraph.add_run(f", {pages}.")

    total -= 1

# 保存文档
doc.save("highlighted_publications.docx")
print("处理完成：highlighted_publications.docx ✅")
