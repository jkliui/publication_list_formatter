import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import logging

# --- 1. Configure Logging ---
# Sets up the logging module for outputting information, warnings, and errors.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- 2. Load Configuration ---
# Loads document and text formatting settings from an external JSON file.
def load_config(config_path="config.json"):
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        logger.info(f"Configuration loaded successfully from {config_path}")
        return config
    except FileNotFoundError:
        logger.error(f"Config file not found at {config_path}. Please create one.")
        raise
    except json.JSONDecodeError:
        logger.error(f"Error decoding JSON from {config_path}. Check file format.")
        raise

config = load_config()

# Retrieve the target author string from the configuration.
target_author = config.get("target_author", "")
if not target_author:
    logger.warning("No 'target_author' specified in config.json. No author will be highlighted.")

# --- 3. Format Authors Function ---
# Processes and formats author names according to specific rules, handling variations with and without commas.
def format_authors(authors_str):
    authors = authors_str.split(';')
    formatted_authors_list = []

    for author_entry in authors:
        author_entry = author_entry.strip()
        if not author_entry:
            continue

        current_initials = []
        last_name_found = ""

        # Handle authors formatted as "Last, First [Middle]"
        if ',' in author_entry:
            parts = author_entry.split(',', 1)
            last_name_found = parts[0].strip()

            if len(parts) > 1:
                first_middle_part = parts[1].strip()
                name_segments = first_middle_part.replace('.', ' ').split()

                for segment in name_segments:
                    if segment:
                        current_initials.append(segment[0].upper())
            
            if current_initials:
                formatted_authors_list.append(
                    f"{last_name_found}, {' '.join([f'{init}.' for init in current_initials])}")
            else:
                formatted_authors_list.append(f"{last_name_found}.")

        else:
            # Handle authors formatted as "First [Middle] Last" or "SingleName"
            name_parts = author_entry.split()
            
            if len(name_parts) >= 2:
                last_name_found = name_parts[-1].strip()
                first_middle_parts = name_parts[:-1]

                for segment in first_middle_parts:
                    if segment:
                        current_initials.append(segment[0].upper())
                
                if current_initials:
                    formatted_authors_list.append(
                        f"{last_name_found}, {' '.join([f'{init}.' for init in current_initials])}")
                else:
                    if len(name_parts) >= 1:
                         formatted_authors_list.append(f"{last_name_found}, {name_parts[0][0].upper()}.")
                    else:
                         formatted_authors_list.append(author_entry.strip())

            else:
                formatted_authors_list.append(author_entry.strip())

    return '; '.join(formatted_authors_list)


# --- 4. Read CSV File ---
# Reads publication data from a CSV file.
csv_file_path = "publication_list.csv"
try:
    df = pd.read_csv(csv_file_path)
    logger.info(f"Loaded {len(df)} rows from {csv_file_path}")
except FileNotFoundError:
    logger.error(f"Input CSV file not found at {csv_file_path}. Please ensure it exists.")
    exit()

# Cleans and converts data types for specified columns in the DataFrame.
df.fillna('', inplace=True)
df['Volume'] = pd.to_numeric(df['Volume'], errors='coerce').fillna(0).astype(int)
df['Year'] = pd.to_numeric(df['Year'], errors='coerce').fillna(0).astype(int)
df = df.astype({col: str for col in df.columns if col not in ['Volume', 'Year']})
logger.info("Data cleaned and types converted.")

# Applies the formatting function to the 'Authors' column.
df['Formatted Authors'] = df['Authors'].apply(format_authors)
logger.info("Authors formatted.")

# Selects relevant columns and sorts the DataFrame by 'Year' in descending order.
df_selected = df[['Formatted Authors', 'Title', 'Publication', 'Year', 'Volume', 'Pages']]
df_selected = df_selected.sort_values(by='Year', ascending=False).reset_index(drop=True)
logger.info("Data selected and sorted by year.")

# --- 5. Create Word Document and Apply Configuration ---
# Initializes a new Word document and applies default font settings from the config.
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = config['document_settings']['default_font_name']
font.size = Pt(config['document_settings']['default_font_size_pt'])
logger.info(f"Default font set to {font.name} at {font.size.pt}pt.")

total_publications = len(df_selected)

# Iterates through each publication entry and adds it to the Word document.
for idx, row in df_selected.iterrows():
    author = row['Formatted Authors']
    title = row['Title']
    journal = row['Publication']
    year = row['Year']
    volume = row['Volume']
    pages = row['Pages']

    paragraph = doc.add_paragraph()
    # Applies paragraph formatting settings from the config.
    para_format = paragraph.paragraph_format
    para_format.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, config['paragraph_formats']['alignment'].upper())
    para_format.line_spacing = config['paragraph_formats']['line_spacing_pt']
    para_format.space_after = Pt(config['paragraph_formats']['space_after_pt'])
    para_format.left_indent = Inches(config['paragraph_formats']['left_indent_inches'])
    para_format.first_line_indent = Inches(config['paragraph_formats']['first_line_indent_inches'])

    # --- 6. Add Runs to Paragraph ---
    # Adds formatted text runs to the current paragraph, applying styles as needed.

    # Writes the publication number (in reverse order).
    paragraph.add_run(f"{total_publications - idx})  ")

    # Writes author names, highlighting the target author if specified.
    if target_author and target_author in author:
        parts = author.split(target_author)
        for i, part in enumerate(parts):
            if part:
                paragraph.add_run(part)
            if i < len(parts) - 1:
                run = paragraph.add_run(target_author)
                if config['text_styles']['target_author_bold']:
                    run.bold = True
                if config['text_styles']['target_author_underline']:
                    run.underline = True
    else:
        paragraph.add_run(author)

    # Adds title, journal, year, volume, and pages with specific formatting.
    paragraph.add_run(f" {title}. ")

    journal_run = paragraph.add_run(journal)
    if config['text_styles']['journal_italic']:
        journal_run.italic = True
    paragraph.add_run(", ")

    year_run = paragraph.add_run(str(year))
    if config['text_styles']['year_bold']:
        year_run.bold = True

    # Adds volume with italic formatting.
    volume_run = paragraph.add_run(f", {volume}")
    if config['text_styles']['volume_italic']:
        volume_run.italic = True

    # Adds pages (if present) and the final period without italic formatting.
    if pages:
        paragraph.add_run(f", {pages}.")
    else:
        paragraph.add_run(".")

    logger.info(f"Processed publication {idx + 1}/{total_publications}: {title[:50]}...")

# Save the document.
doc_output_path = "highlighted_publications.docx"
try:
    doc.save(doc_output_path)
    logger.info(f"Processing complete: {doc_output_path} ")
except Exception as e:
    logger.error(f"Error saving document to {doc_output_path}: {e}")